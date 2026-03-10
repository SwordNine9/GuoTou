import json
import logging
import os
import re
import tempfile
from io import BytesIO
from typing import Any

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openai import OpenAI

matplotlib.use("Agg")

LOG_DIR = "logs"
LOG_FILE = os.path.join(LOG_DIR, "app.log")
SECTION_CONFIG_FILE = "section_mapping.json"

DEFAULT_SECTION_CONFIG = {
    "chart_anchor": "（一）全国电力供应数据",
    "sections": [
        {"title": "二、政府文件", "json_path": "gov_policies", "content_type": "policy"},
        {"title": "一、电力交易新政", "json_path": "energy_new_policies.电力交易新政", "content_type": "energy"},
        {"title": "二、区域电价政策", "json_path": "energy_new_policies.区域电价政策", "content_type": "energy"},
        {"title": "三、重点开发政策", "json_path": "energy_new_policies.重点开发政策", "content_type": "energy"},
        {"title": "参考资料", "json_path": "references", "content_type": "reference"},
    ],
}


def setup_logger():
    os.makedirs(LOG_DIR, exist_ok=True)
    logger = logging.getLogger("guotou_service")
    if logger.handlers:
        return logger
    logger.setLevel(logging.INFO)
    handler = logging.FileHandler(LOG_FILE, encoding="utf-8")
    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    return logger


LOGGER = setup_logger()


def set_run_font(run, font_name="仿宋", size=None, is_bold=False):
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    r.rPr.rFonts.set(qn("w:ascii"), font_name)
    r.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size:
        run.font.size = size
    if is_bold:
        run.bold = True


def tokenize_for_rag(text):
    if not text:
        return []
    return re.findall(r"[\u4e00-\u9fff]+|[a-zA-Z0-9_]+", text.lower())


def split_text_into_chunks(text, chunk_size=260, overlap=40):
    cleaned = " ".join((text or "").split())
    if not cleaned:
        return []
    chunks = []
    step = max(1, chunk_size - overlap)
    for start in range(0, len(cleaned), step):
        chunks.append(cleaned[start : start + chunk_size])
    return chunks


def build_rag_chunks_from_texts(file_texts: list[dict[str, str]]):
    rag_chunks = []
    for item in file_texts or []:
        source = item.get("name", "unknown")
        text = item.get("text", "")
        for chunk in split_text_into_chunks(text):
            rag_chunks.append({"source": source, "content": chunk})
    return rag_chunks


def retrieve_rag_context(query, rag_chunks, top_k=3):
    if not rag_chunks:
        return ""
    query_tokens = set(tokenize_for_rag(query))
    if not query_tokens:
        return ""

    scored = []
    for chunk in rag_chunks:
        chunk_tokens = set(tokenize_for_rag(chunk["content"]))
        overlap = len(query_tokens & chunk_tokens)
        if overlap > 0:
            scored.append((overlap, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    return "\n".join([f"[来源: {item['source']}] {item['content']}" for _, item in scored[:top_k]])


def generate_policy_analysis(client, summary, rag_context=""):
    if not client:
        return "（因未配置 API Key，系统跳过 AI 解析生成）"

    user_prompt = f"政策摘要：{summary}"
    if rag_context:
        user_prompt += f"\n\n以下是可参考内部资料（RAG 检索结果），请优先依据资料，不要虚构：\n{rag_context}"

    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {
                    "role": "system",
                    "content": "你是一位电力政策分析专家。请基于摘要给出100字左右的影响分析和应对建议。若提供参考资料，优先依据资料并保持可追溯。",
                },
                {"role": "user", "content": user_prompt},
            ],
            stream=False,
        )
        return response.choices[0].message.content
    except Exception as err:
        LOGGER.exception("AI解析失败")
        return f"AI 解析生成失败: {str(err)}"


def load_section_config(config_path=SECTION_CONFIG_FILE):
    if not os.path.exists(config_path):
        return DEFAULT_SECTION_CONFIG
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        LOGGER.exception("section mapping 加载失败，回退默认配置")
        return DEFAULT_SECTION_CONFIG


def get_by_json_path(data, path):
    cur = data
    for key in path.split("."):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(key)
    return cur


def generate_pie_chart(imp_data):
    categories = ["Hydro", "Thermal", "Nuclear", "Solar", "Wind"]
    keys = ["hydro_capacity", "thermal_capacity", "nuclear_capacity", "solar_capacity", "wind_capacity"]
    values = []
    for k in keys:
        try:
            values.append(float(imp_data.get(k, 0) or 0))
        except ValueError:
            values.append(0)

    if sum(values) == 0:
        return None

    tfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.figure(figsize=(6, 6))
    plt.pie(values, labels=categories, autopct="%.1f%%", startangle=90)
    plt.title("Power Installed Capacity Structure")
    plt.tight_layout()
    plt.savefig(tfile.name)
    plt.close()
    return tfile.name


def validate_json_data(json_data):
    errors = []
    if not isinstance(json_data, dict):
        return ["JSON 顶层必须是对象"]

    required_top_fields = ["important_data_values", "gov_policies", "energy_new_policies", "references"]
    for field in required_top_fields:
        if field not in json_data:
            errors.append(f"缺少顶层字段：{field}")

    imp = json_data.get("important_data_values", {})
    for k in ["report_date", "total_capacity", "hydro_capacity", "thermal_capacity"]:
        if not imp.get(k):
            errors.append(f"important_data_values.{k} 为空")

    return errors


def validate_template(doc, config, expected_placeholder_count=31):
    full_text = "\n".join([p.text for p in doc.paragraphs])
    missing_titles = [s["title"] for s in config.get("sections", []) if s["title"] not in full_text]

    chart_anchor = config.get("chart_anchor")
    chart_anchor_ok = chart_anchor in full_text
    placeholder_count = full_text.count("**")
    placeholder_ok = placeholder_count >= expected_placeholder_count

    return {
        "missing_titles": missing_titles,
        "chart_anchor_ok": chart_anchor_ok,
        "placeholder_count": placeholder_count,
        "placeholder_ok": placeholder_ok,
        "expected_placeholder_count": expected_placeholder_count,
    }


def insert_content_after_keyword(
    doc,
    keyword,
    content_data,
    content_type,
    client,
    rag_chunks=None,
    rag_top_k=3,
    include_rag_snippets=False,
):
    target_index = -1
    for i, paragraph in enumerate(doc.paragraphs):
        if keyword in paragraph.text:
            target_index = i
            break

    if target_index == -1:
        return False

    base_node = doc.add_paragraph("") if target_index == len(doc.paragraphs) - 1 else doc.paragraphs[target_index + 1]
    font_size_standard = Pt(16)

    for item in content_data or []:
        if content_type == "reference":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"{item.get('source', '')}：{item.get('title', '')} ({item.get('url', '')})")
            set_run_font(run, font_name="仿宋", size=font_size_standard)
            continue

        p = base_node.insert_paragraph_before()
        label = "政策标题" if content_type == "energy" else "政策名称"
        run = p.add_run(f"{label}：{item.get('title', '')}")
        set_run_font(run, font_name="仿宋", size=font_size_standard, is_bold=True)

        if content_type == "energy":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"区域：{item.get('region', '')}")
            set_run_font(run, font_name="仿宋", size=font_size_standard, is_bold=True)

        if content_type == "policy":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"发布机构：{item.get('agency', '')}")
            set_run_font(run, font_name="仿宋", size=font_size_standard)

        p = base_node.insert_paragraph_before()
        date_label = "发布日期" if content_type == "energy" else "发布时间"
        run = p.add_run(f"{date_label}：{item.get('date', '')}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        summary = item.get("summary", "")
        p = base_node.insert_paragraph_before()
        run = p.add_run(f"政策摘要：{summary}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        query = f"{item.get('title', '')} {summary}"
        rag_context = retrieve_rag_context(query, rag_chunks or [], top_k=rag_top_k)
        analysis = generate_policy_analysis(client, summary, rag_context=rag_context)

        p = base_node.insert_paragraph_before()
        run = p.add_run(f"政策解析：{analysis}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        if include_rag_snippets and rag_context:
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"RAG参考片段：{rag_context}")
            set_run_font(run, font_name="仿宋", size=Pt(12))

        base_node.insert_paragraph_before("")
    return True


def process_document(
    template_file,
    json_data,
    api_key="",
    section_config=None,
    rag_chunks=None,
    rag_top_k=3,
    include_rag_snippets=False,
):
    section_config = section_config or DEFAULT_SECTION_CONFIG
    doc = Document(template_file)
    imp_data = json_data.get("important_data_values", {})

    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com") if api_key else None

    fill_order = [
        imp_data.get("report_date"),
        imp_data.get("total_capacity"), imp_data.get("total_capacity_yoy"),
        imp_data.get("hydro_capacity"), imp_data.get("hydro_yoy"),
        imp_data.get("thermal_capacity"), imp_data.get("thermal_yoy"),
        imp_data.get("nuclear_capacity"), imp_data.get("nuclear_yoy"),
        imp_data.get("solar_capacity"), imp_data.get("solar_yoy"),
        imp_data.get("wind_capacity"), imp_data.get("wind_yoy"),
        imp_data.get("report_date"),
        imp_data.get("market_trade"), imp_data.get("market_trade_yoy"),
        imp_data.get("provincial_trade"), imp_data.get("provincial_trade_yoy"),
        imp_data.get("cross_region_trade"), imp_data.get("cross_region_trade_yoy"),
        imp_data.get("green_trade"), imp_data.get("green_trade_yoy"),
        imp_data.get("report_date"),
        imp_data.get("society_usage"), imp_data.get("society_usage_yoy"),
        imp_data.get("primary_usage"), imp_data.get("primary_usage_yoy"),
        imp_data.get("secondary_usage"), imp_data.get("secondary_usage_yoy"),
        imp_data.get("tertiary_usage"), imp_data.get("tertiary_usage_yoy"),
    ]

    values_iter = iter(fill_order)
    for paragraph in doc.paragraphs:
        if "**" in paragraph.text:
            for _ in range(paragraph.text.count("**")):
                try:
                    val = next(values_iter)
                    for run in paragraph.runs:
                        if "**" in run.text:
                            run.text = run.text.replace("**", str(val), 1)
                            set_run_font(run, font_name="仿宋")
                            break
                except StopIteration:
                    break

    chart_path = generate_pie_chart(imp_data)
    if chart_path:
        chart_anchor = section_config.get("chart_anchor", "（一）全国电力供应数据")
        for i, p in enumerate(doc.paragraphs):
            if chart_anchor in p.text and i + 1 < len(doc.paragraphs):
                target_p = doc.paragraphs[i + 1]
                run = target_p.add_run()
                run.add_break()
                run.add_picture(chart_path, width=Inches(5.0))
                break
        os.unlink(chart_path)

    not_found_sections = []
    for section in section_config.get("sections", []):
        section_data = get_by_json_path(json_data, section["json_path"]) or []
        ok = insert_content_after_keyword(
            doc,
            section["title"],
            section_data,
            section["content_type"],
            client,
            rag_chunks=rag_chunks,
            rag_top_k=rag_top_k,
            include_rag_snippets=include_rag_snippets,
        )
        if not ok:
            not_found_sections.append(section["title"])

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, {"not_found_sections": not_found_sections}


def validate_payload(json_data: dict[str, Any], template_bytes: bytes, section_config=None):
    section_config = section_config or DEFAULT_SECTION_CONFIG
    data_errors = validate_json_data(json_data)
    doc = Document(BytesIO(template_bytes))
    template_result = validate_template(doc, section_config)
    return {"data_errors": data_errors, "template_check": template_result}
