import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from openai import OpenAI

# ========== 配置部分 ==========
DEEPSEEK_API_URL = "https://api.deepseek.com"
DEEPSEEK_API_KEY = "YOUR_DEEPSEEK_API_KEY" # 请替换

template_path = "国投电力业务发展部政策资讯专刊（模板）.docx"
output_docx = "政策资讯专刊_自动生成.docx"
# 图表路径
chart_path = "装机容量结构饼图.png"

# ========== 辅助函数：设置中文字体 ==========
def set_run_font(run, font_name='仿宋_GB2312', size=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = size

# ========== 核心功能：在指定关键词段落后插入内容 ==========
def insert_content_after_keyword(doc, keyword, content_data, content_type="policy"):
    """
    keyword: 模板中的标题文本（如"二、政府文件"）
    content_data: 要插入的数据列表
    content_type: 'policy' (政府文件), 'energy' (能源新政), 'reference' (参考资料)
    """
    target_index = -1
    # 1. 找到关键词所在的段落索引
    for i, paragraph in enumerate(doc.paragraphs):
        if keyword in paragraph.text:
            target_index = i
            break
    
    if target_index == -1:
        print(f"未找到章节：{keyword}，跳过插入。")
        return

    # 2. 倒序插入，保证顺序正确（因为insert_paragraph_before是在当前段落前插入）
    # 我们需要在 target_index + 1 的位置开始插入，为了保持逻辑简单，
    # 我们找到下一段，然后在它前面不断插入。
    
    # 如果是最后一段，直接append
    if target_index == len(doc.paragraphs) - 1:
        insert_anchor = None 
    else:
        insert_anchor = doc.paragraphs[target_index + 1]

    # 辅助插入函数
    def add_p(text, bold=False, style=None):
        if insert_anchor:
            p = insert_anchor.insert_paragraph_before(text, style=style)
        else:
            p = doc.add_paragraph(text, style=style)
        
        # 简单格式化
        if bold:
            p.runs[0].bold = True
        # 设置统一字体，避免格式突兀
        if p.runs:
            set_run_font(p.runs[0], '仿宋_GB2312', Pt(12))
        return p

    # 根据不同类型处理数据插入
    if content_type == "policy":
        for policy in reversed(content_data): # 倒序遍历，配合 insert_before 逻辑可能需要调整，这里简化处理
            # 注意：insert_paragraph_before 会把新段落挤在 anchor 前面。
            # 为了让顺序正常（1, 2, 3），我们需要倒序插入列表，或者改变插入策略。
            # 这里使用简单的追加逻辑修正：直接在找到的段落后插入不太容易，
            # 最稳妥的方法是：在 Python-docx 中操作 XML，或者利用 insert_paragraph_before
            pass 
            
    # === 修正后的插入逻辑：直接操作列表 ===
    # 为了简化，我们采用不依赖 insert_anchor 的复杂逻辑，而是直接在特定段落后方插入
    # 但 python-docx 原生不支持 insert_after。
    # 变通方案：我们在模板里必须预留好空行，或者使用下面的 move_table_after 逻辑(太复杂)。
    # 
    # 最简单的方案：
    # 我们假设模板里，标题下面紧接着就是正文。
    # 我们反向遍历数据，使用 insert_paragraph_before 插入到 "target_index + 1" 的段落前面。
    
    if insert_anchor:
        base_node = insert_anchor
    else:
        # 如果标题是最后一行，先加个空段落作为锚点
        base_node = doc.add_paragraph("")

    # 倒序处理数据，这样再次 insert_before 时顺序就正过来了
    for item in reversed(content_data):
        # 插入空行分隔
        base_node.insert_paragraph_before("") 
        
        if content_type == "policy":
            # 插入解析
            analysis = generate_policy_analysis(item['summary'])
            p = base_node.insert_paragraph_before(f"政策解析：{analysis}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"政策摘要：{item['summary']}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"发布时间：{item['date']}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"发布机构：{item['agency']}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"政策名称：{item['title']}")
            p.runs[0].bold = True
            set_run_font(p.runs[0], size=Pt(12))

        elif content_type == "energy":
            # item 结构: {'region':..., 'title':...}
            analysis = generate_policy_analysis(item['summary'])
            
            p = base_node.insert_paragraph_before(f"政策解析：{analysis}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"政策摘要：{item['summary']}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"发布日期：{item['date']}")
            set_run_font(p.runs[0])

            p = base_node.insert_paragraph_before(f"政策标题：{item['title']}")
            set_run_font(p.runs[0])
            
            p = base_node.insert_paragraph_before(f"区域：{item['region']}")
            p.runs[0].bold = True
            set_run_font(p.runs[0])

        elif content_type == "reference":
             # item 结构: {'source':..., 'title':..., 'url':...}
            p = base_node.insert_paragraph_before(f"{item['source']}：{item['title']} ({item['url']})")
            set_run_font(p.runs[0], '楷体_GB2312', Pt(11))

# ========== DeepSeek API (保持不变，建议增加错误处理) ==========
def generate_policy_analysis(summary):
    # 为了演示，如果没Key直接返回模拟数据
    if "YOUR_DEEPSEEK_API_KEY" in DEEPSEEK_API_KEY:
        return "（此处为自动生成的AI政策解析，因未配置API Key显示测试文本）"
    
    try:
        client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_API_URL)
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是一位电力政策分析专家，请简要分析以下政策的影响和建议。"},
                {"role": "user", "content": summary}
            ],
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"AI分析生成失败: {e}")
        return "暂无解析"

# ========== 图表生成 (保持不变) ==========
def generate_capacity_pie_chart(important_values, output_path):
    # ... (你的原始代码逻辑保持不变) ...
    categories = ["Hydro", "Thermal", "Nuclear", "Solar", "Wind"]
    keys = ["hydro_capacity", "thermal_capacity", "nuclear_capacity", "solar_capacity", "wind_capacity"]
    values = [float(important_values.get(k, 0) or 0) for k in keys]
    
    if sum(values) == 0: return

    plt.figure(figsize=(6, 6))
    plt.pie(values, labels=categories, autopct="%.1f%%", startangle=90)
    plt.title("Power Installed Capacity Structure")
    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()

# ========== 主程序逻辑 ==========
def main():
    # 1. 加载数据
    with open("policy_data.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    
    imp_data = data.get("important_data_values", {})
    
    # 2. 加载模板
    try:
        doc = Document(template_path)
    except:
        print(f"找不到模板文件：{template_path}")
        return

    # 3. 填充 "**" 占位符 (保持你的逻辑，但建议加个Try保护)
    # 注意：这里完全依赖模板中 ** 出现的顺序，这很脆弱。
    # 如果模板变了，这里必须手动调整顺序。
    fill_order = [
        imp_data.get("report_date"),
        imp_data.get("total_capacity"), imp_data.get("total_capacity_yoy"),
        imp_data.get("hydro_capacity"), imp_data.get("hydro_yoy"),
        imp_data.get("thermal_capacity"), imp_data.get("thermal_yoy"),
        imp_data.get("nuclear_capacity"), imp_data.get("nuclear_yoy"),
        imp_data.get("solar_capacity"), imp_data.get("solar_yoy"),
        imp_data.get("wind_capacity"), imp_data.get("wind_yoy"),
        imp_data.get("report_date"), # 交易部分日期
        imp_data.get("market_trade"), imp_data.get("market_trade_yoy"),
        imp_data.get("provincial_trade"), imp_data.get("provincial_trade_yoy"),
        imp_data.get("cross_region_trade"), imp_data.get("cross_region_trade_yoy"),
        imp_data.get("green_trade"), imp_data.get("green_trade_yoy"),
        imp_data.get("report_date"), # 用电部分日期
        imp_data.get("society_usage"), imp_data.get("society_usage_yoy"),
        imp_data.get("primary_usage"), imp_data.get("primary_usage_yoy"),
        imp_data.get("secondary_usage"), imp_data.get("secondary_usage_yoy"),
        imp_data.get("tertiary_usage"), imp_data.get("tertiary_usage_yoy")
    ]
    
    values_iter = iter(fill_order)
    
    for paragraph in doc.paragraphs:
        # 特殊处理：在"全国电力供应数据"这段里插入图片
        # 这是一个简单的定位方式
        if "（一）全国电力供应数据" in paragraph.text:
            # 在此段落后占个位，稍后插入图片
            # 由于 text replacement 不容易插入图片，我们记录这个 paragraph
            pass 

        if "**" in paragraph.text:
            # 简单的全替换逻辑，处理单段落多个占位符
            count = paragraph.text.count("**")
            for _ in range(count):
                try:
                    val = next(values_iter)
                    # 替换逻辑：这里用简单replace，不保留原格式可能会丢bold
                    # 你的原始函数 replace_next_placeholder 更好，这里直接内联简化版
                    for run in paragraph.runs:
                        if "**" in run.text:
                            run.text = run.text.replace("**", str(val), 1)
                            break
                except StopIteration:
                    break

    # 4. 生成饼图并插入 (插在 "（一）全国电力供应数据" 后面)
    generate_capacity_pie_chart(imp_data, chart_path)
    
    # 寻找插入位置
    for i, p in enumerate(doc.paragraphs):
        if "（一）全国电力供应数据" in p.text:
            # 在标题的下一段（通常是数据段）之后插入图表
            # 这里选择在数据段(也就是包含"截至**"的那一段)之后插入
            target_p = doc.paragraphs[i+1] 
            run = target_p.add_run()
            run.add_break()
            run.add_picture(chart_path, width=Inches(5.0))
            break

    # 5. 插入政府文件 (定位到 "二、政府文件")
    # 模板中有 "二、政府文件 （墨体三号加粗）"，我们匹配关键字
    insert_content_after_keyword(doc, "二、政府文件", data.get("gov_policies", []), "policy")

    # 6. 插入能源新政
    # 这里的逻辑比较复杂，因为JSON里是分了"电力交易新政"、"区域电价"等子类
    # 而模板里有 "一、电力交易新政" 等子标题。我们需要分别匹配。
    new_policies = data.get("energy_new_policies", {})
    
    # 匹配 "一、电力交易新政"
    if "电力交易新政" in new_policies:
        insert_content_after_keyword(doc, "一、电力交易新政", new_policies["电力交易新政"], "energy")
        
    # 匹配 "二、区域电价政策"
    if "区域电价政策" in new_policies:
        insert_content_after_keyword(doc, "二、区域电价政策", new_policies["区域电价政策"], "energy")
        
    # 匹配 "三、重点开发政策"
    if "重点开发政策" in new_policies:
        insert_content_after_keyword(doc, "三、重点开发政策", new_policies["重点开发政策"], "energy")

    # 7. 插入参考资料 (定位到 "他山之石")
    # 模板中是 "他山之石"，下面有 "参考资料"
    insert_content_after_keyword(doc, "参考资料", data.get("references", []), "reference")

    # 8. 保存
    doc.save(output_docx)
    print(f"文档生成完毕：{output_docx}")

if __name__ == "__main__":
    main()