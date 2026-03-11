import json
import logging
import math
import os
import re
import tempfile
import time
from datetime import datetime
from io import BytesIO

import matplotlib
import matplotlib.pyplot as plt
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openai import OpenAI

# 设置 matplotlib 后端，防止在无窗口环境下报错
matplotlib.use("Agg")

LOG_DIR = "logs"
LOG_FILE = os.path.join(LOG_DIR, "app.log")
HISTORY_FILE = "generation_history.jsonl"
SECTION_CONFIG_FILE = "section_mapping.json"

DEFAULT_SECTION_CONFIG = {
    "chart_anchor": "（一）全国电力供应数据",
    "sections": [
        {"title": "二、政府文件", "json_path": "gov_policies", "content_type": "policy"},
        {
            "title": "一、电力交易新政",
            "json_path": "energy_new_policies.电力交易新政",
            "content_type": "energy",
        },
        {
            "title": "二、区域电价政策",
            "json_path": "energy_new_policies.区域电价政策",
            "content_type": "energy",
        },
        {
            "title": "三、重点开发政策",
            "json_path": "energy_new_policies.重点开发政策",
            "content_type": "energy",
        },
        {"title": "参考资料", "json_path": "references", "content_type": "reference"},
    ],
}

# ================= 1. 页面基础配置 =================
st.set_page_config(page_title="电力政策专刊生成助手", page_icon="⚡", layout="wide")

st.markdown(
    """
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button {
        width: 100%;
        background-color: #0066cc;
        color: white;
        font-weight: bold;
        height: 50px;
        border-radius: 8px;
    }
    .stButton>button:hover {
        background-color: #0052a3;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def setup_logger():
    os.makedirs(LOG_DIR, exist_ok=True)
    logger = logging.getLogger("guotou_app")
    if logger.handlers:
        return logger

    logger.setLevel(logging.INFO)
    handler = logging.FileHandler(LOG_FILE, encoding="utf-8")
    formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    return logger


LOGGER = setup_logger()


# ================= 2. 工具函数 =================
def set_run_font(run, font_name="仿宋", size=None, is_bold=False):
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    r.rPr.rFonts.set(qn("w:ascii"), font_name)
    r.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size:
        run.font.size = size
    if is_bold:
        run.bold = True


def tokenize_for_rag(text):
    if not text:
        return []
    return re.findall(r"[\u4e00-\u9fff]+|[a-zA-Z0-9_]+", text.lower())


def split_text_into_chunks(text, chunk_size=260, overlap=40):
    cleaned = " ".join((text or "").split())
    if not cleaned:
        return []

    chunks = []
    step = max(1, chunk_size - overlap)
    for start in range(0, len(cleaned), step):
        chunks.append(cleaned[start : start + chunk_size])
    return chunks


def build_rag_chunks(uploaded_files):
    rag_chunks = []
    for file in uploaded_files or []:
        try:
            text = file.getvalue().decode("utf-8")
        except UnicodeDecodeError:
            text = file.getvalue().decode("utf-8", errors="ignore")

        for chunk in split_text_into_chunks(text):
            rag_chunks.append({"source": file.name, "content": chunk})
    return rag_chunks


def build_vector_index(rag_chunks):
    if not rag_chunks:
        return {"idf": {}, "chunks": []}

    doc_freq = {}
    tokenized_chunks = []
    for chunk in rag_chunks:
        tokens = tokenize_for_rag(chunk.get("content", ""))
        tokenized_chunks.append(tokens)
        for token in set(tokens):
            doc_freq[token] = doc_freq.get(token, 0) + 1

    n_docs = len(rag_chunks)
    idf = {tok: math.log((n_docs + 1) / (df + 1)) + 1 for tok, df in doc_freq.items()}

    indexed_chunks = []
    for chunk, tokens in zip(rag_chunks, tokenized_chunks):
        tf = {}
        for token in tokens:
            tf[token] = tf.get(token, 0) + 1
        total = max(1, len(tokens))

        vec = {}
        for token, count in tf.items():
            vec[token] = (count / total) * idf.get(token, 0.0)

        norm = math.sqrt(sum(v * v for v in vec.values()))
        indexed_chunks.append({**chunk, "vec": vec, "norm": norm})

    return {"idf": idf, "chunks": indexed_chunks}


def retrieve_rag_context(query, rag_chunks, top_k=3):
    if not rag_chunks:
        return ""

    vector_index = build_vector_index(rag_chunks)
    chunks = vector_index.get("chunks", [])
    idf = vector_index.get("idf", {})

    query_tokens = tokenize_for_rag(query)
    if not query_tokens:
        return ""

    q_tf = {}
    for token in query_tokens:
        q_tf[token] = q_tf.get(token, 0) + 1

    q_total = max(1, len(query_tokens))
    q_vec = {t: (c / q_total) * idf.get(t, 0.0) for t, c in q_tf.items()}
    q_norm = math.sqrt(sum(v * v for v in q_vec.values()))
    if q_norm == 0:
        return ""

    scored = []
    for chunk in chunks:
        dot = 0.0
        for token, q_val in q_vec.items():
            dot += q_val * chunk["vec"].get(token, 0.0)
        if dot <= 0 or chunk["norm"] == 0:
            continue
        sim = dot / (q_norm * chunk["norm"])
        scored.append((sim, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    return "\n".join([f"[来源: {item['source']}] {item['content']}" for _, item in scored[:top_k]])


def generate_policy_analysis(client, summary, rag_context=""):
    if not client:
        return "（因未配置 API Key，系统跳过 AI 解析生成）"

    user_prompt = f"政策摘要：{summary}"
    if rag_context:
        user_prompt += (
            "\n\n以下是可参考内部资料（RAG 检索结果），请优先依据资料，不要虚构：\n"
            f"{rag_context}"
        )

    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {
                    "role": "system",
                    "content": "你是一位电力政策分析专家。请基于摘要给出100字左右的影响分析和应对建议。若提供参考资料，优先依据资料并保持可追溯。",
                },
                {"role": "user", "content": user_prompt},
            ],
            stream=False,
        )
        return response.choices[0].message.content
    except Exception as err:
        LOGGER.exception("AI解析失败")
        return f"AI 解析生成失败: {str(err)}"


def load_section_config():
    if not os.path.exists(SECTION_CONFIG_FILE):
        return DEFAULT_SECTION_CONFIG
    try:
        with open(SECTION_CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        LOGGER.exception("section_mapping.json 加载失败，回退默认配置")
        return DEFAULT_SECTION_CONFIG


def get_by_json_path(data, path):
    cur = data
    for key in path.split("."):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(key)
    return cur


def generate_pie_chart(imp_data):
    categories = ["Hydro", "Thermal", "Nuclear", "Solar", "Wind"]
    keys = [
        "hydro_capacity",
        "thermal_capacity",
        "nuclear_capacity",
        "solar_capacity",
        "wind_capacity",
    ]

    values = []
    for k in keys:
        try:
            values.append(float(imp_data.get(k, 0) or 0))
        except ValueError:
            values.append(0)

    if sum(values) == 0:
        return None

    tfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.figure(figsize=(6, 6))
    plt.pie(values, labels=categories, autopct="%.1f%%", startangle=90)
    plt.title("Power Installed Capacity Structure")
    plt.tight_layout()
    plt.savefig(tfile.name)
    plt.close()
    return tfile.name


def validate_json_data(json_data):
    errors = []
    if not isinstance(json_data, dict):
        return ["JSON 顶层必须是对象"]

    required_top_fields = ["important_data_values", "gov_policies", "energy_new_policies", "references"]
    for field in required_top_fields:
        if field not in json_data:
            errors.append(f"缺少顶层字段：{field}")

    imp = json_data.get("important_data_values", {})
    for k in ["report_date", "total_capacity", "hydro_capacity", "thermal_capacity"]:
        if not imp.get(k):
            errors.append(f"important_data_values.{k} 为空")

    return errors


def validate_template(doc, config, expected_placeholder_count):
    full_text = "\n".join([p.text for p in doc.paragraphs])
    missing_titles = []
    for section in config.get("sections", []):
        if section["title"] not in full_text:
            missing_titles.append(section["title"])

    chart_anchor = config.get("chart_anchor")
    chart_anchor_ok = chart_anchor in full_text

    placeholder_count = full_text.count("**")
    placeholder_ok = placeholder_count >= expected_placeholder_count

    return {
        "missing_titles": missing_titles,
        "chart_anchor_ok": chart_anchor_ok,
        "placeholder_count": placeholder_count,
        "placeholder_ok": placeholder_ok,
        "expected_placeholder_count": expected_placeholder_count,
    }


def append_history(record):
    with open(HISTORY_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


def read_recent_history(limit=8):
    if not os.path.exists(HISTORY_FILE):
        return []
    with open(HISTORY_FILE, "r", encoding="utf-8") as f:
        lines = [line.strip() for line in f if line.strip()]
    return [json.loads(x) for x in lines[-limit:]][::-1]


# ================= 3. 文档处理核心逻辑 =================
def insert_content_after_keyword(
    doc,
    keyword,
    content_data,
    content_type,
    client,
    rag_chunks=None,
    rag_top_k=3,
    include_rag_snippets=False,
):
    target_index = -1
    for i, paragraph in enumerate(doc.paragraphs):
        if keyword in paragraph.text:
            target_index = i
            break

    if target_index == -1:
        st.warning(f"⚠️ 模板中未找到章节标题：“{keyword}”，已跳过该部分。")
        return

    base_node = doc.add_paragraph("") if target_index == len(doc.paragraphs) - 1 else doc.paragraphs[target_index + 1]
    font_size_standard = Pt(16)

    for item in content_data or []:
        if content_type == "reference":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"{item.get('source', '')}：{item.get('title', '')} ({item.get('url', '')})")
            set_run_font(run, font_name="仿宋", size=font_size_standard)
            continue

        p = base_node.insert_paragraph_before()
        label = "政策标题" if content_type == "energy" else "政策名称"
        run = p.add_run(f"{label}：{item.get('title', '')}")
        set_run_font(run, font_name="仿宋", size=font_size_standard, is_bold=True)

        if content_type == "energy":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"区域：{item.get('region', '')}")
            set_run_font(run, font_name="仿宋", size=font_size_standard, is_bold=True)

        if content_type == "policy":
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"发布机构：{item.get('agency', '')}")
            set_run_font(run, font_name="仿宋", size=font_size_standard)

        p = base_node.insert_paragraph_before()
        date_label = "发布日期" if content_type == "energy" else "发布时间"
        run = p.add_run(f"{date_label}：{item.get('date', '')}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        summary = item.get("summary", "")
        p = base_node.insert_paragraph_before()
        run = p.add_run(f"政策摘要：{summary}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        query = f"{item.get('title', '')} {summary}"
        rag_context = retrieve_rag_context(query, rag_chunks or [], top_k=rag_top_k)
        analysis = generate_policy_analysis(client, summary, rag_context=rag_context)

        p = base_node.insert_paragraph_before()
        run = p.add_run(f"政策解析：{analysis}")
        set_run_font(run, font_name="仿宋", size=font_size_standard)

        if include_rag_snippets and rag_context:
            p = base_node.insert_paragraph_before()
            run = p.add_run(f"RAG参考片段：{rag_context}")
            set_run_font(run, font_name="仿宋", size=Pt(12))

        base_node.insert_paragraph_before("")


def process_document(
    template_file,
    json_data,
    api_key,
    section_config,
    rag_chunks=None,
    rag_top_k=3,
    include_rag_snippets=False,
):
    doc = Document(template_file)
    imp_data = json_data.get("important_data_values", {})

    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com") if api_key else None

    fill_order = [
        imp_data.get("report_date"),
        imp_data.get("total_capacity"), imp_data.get("total_capacity_yoy"),
        imp_data.get("hydro_capacity"), imp_data.get("hydro_yoy"),
        imp_data.get("thermal_capacity"), imp_data.get("thermal_yoy"),
        imp_data.get("nuclear_capacity"), imp_data.get("nuclear_yoy"),
        imp_data.get("solar_capacity"), imp_data.get("solar_yoy"),
        imp_data.get("wind_capacity"), imp_data.get("wind_yoy"),
        imp_data.get("report_date"),
        imp_data.get("market_trade"), imp_data.get("market_trade_yoy"),
        imp_data.get("provincial_trade"), imp_data.get("provincial_trade_yoy"),
        imp_data.get("cross_region_trade"), imp_data.get("cross_region_trade_yoy"),
        imp_data.get("green_trade"), imp_data.get("green_trade_yoy"),
        imp_data.get("report_date"),
        imp_data.get("society_usage"), imp_data.get("society_usage_yoy"),
        imp_data.get("primary_usage"), imp_data.get("primary_usage_yoy"),
        imp_data.get("secondary_usage"), imp_data.get("secondary_usage_yoy"),
        imp_data.get("tertiary_usage"), imp_data.get("tertiary_usage_yoy"),
    ]

    values_iter = iter(fill_order)
    for paragraph in doc.paragraphs:
        if "**" in paragraph.text:
            for _ in range(paragraph.text.count("**")):
                try:
                    val = next(values_iter)
                    for run in paragraph.runs:
                        if "**" in run.text:
                            run.text = run.text.replace("**", str(val), 1)
                            set_run_font(run, font_name="仿宋")
                            break
                except StopIteration:
                    break

    chart_path = generate_pie_chart(imp_data)
    if chart_path:
        chart_anchor = section_config.get("chart_anchor", "（一）全国电力供应数据")
        for i, p in enumerate(doc.paragraphs):
            if chart_anchor in p.text and i + 1 < len(doc.paragraphs):
                target_p = doc.paragraphs[i + 1]
                run = target_p.add_run()
                run.add_break()
                run.add_picture(chart_path, width=Inches(5.0))
                break
        os.unlink(chart_path)

    for section in section_config.get("sections", []):
        section_data = get_by_json_path(json_data, section["json_path"]) or []
        insert_content_after_keyword(
            doc,
            section["title"],
            section_data,
            section["content_type"],
            client,
            rag_chunks=rag_chunks,
            rag_top_k=rag_top_k,
            include_rag_snippets=include_rag_snippets,
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ================= 4. Streamlit 前端界面 =================
st.title("国投电力政策资讯专刊生成系统")
section_config = load_section_config()

with st.sidebar:
    st.header("⚙️ 参数设置")
    api_key = st.text_input("DeepSeek API Key", type="password", help="输入 Key 以启用 AI 自动解析功能")

    st.subheader("📚 RAG 设置")
    enable_rag = st.toggle("启用 RAG 增强解析", value=False)
    rag_top_k = st.slider("RAG 检索片段数", min_value=1, max_value=8, value=3)
    include_rag_snippets = st.checkbox("在报告中写入RAG参考片段", value=False)
    rag_files = st.file_uploader(
        "上传知识库文件（txt/md）", type=["txt", "md"], accept_multiple_files=True
    )

    st.subheader("🧩 模板映射")
    st.caption(f"映射文件：{SECTION_CONFIG_FILE}（不存在时自动使用默认配置）")
    if st.checkbox("显示当前章节映射配置", value=False):
        st.json(section_config)

    st.info("💡 模板预检查会在生成前执行，提示标题缺失与占位符异常。")

rag_chunks = build_rag_chunks(rag_files) if enable_rag else []
if enable_rag:
    st.sidebar.caption(f"已加载 RAG 文本块：{len(rag_chunks)}")

with st.sidebar.expander("🕘 最近生成历史", expanded=False):
    for row in read_recent_history():
        st.write(f"{row.get('timestamp')} | {row.get('status')} | {row.get('duration_s')}s")

history_rows = read_recent_history(limit=50)
if history_rows:
    success_count = len([x for x in history_rows if x.get("status") == "success"])
    success_rate = round(success_count / len(history_rows) * 100, 1)
    avg_duration = round(sum(float(x.get("duration_s", 0)) for x in history_rows) / len(history_rows), 2)
else:
    success_rate = 0.0
    avg_duration = 0.0

st.subheader("📊 运行指标看板")
m1, m2, m3, m4 = st.columns(4)
m1.metric("最近任务数", len(history_rows))
m2.metric("成功率", f"{success_rate}%")
m3.metric("平均耗时", f"{avg_duration}s")
m4.metric("当前RAG文本块", len(rag_chunks))

col1, col2 = st.columns(2)
with col1:
    st.subheader("📂 1. 上传数据 (JSON)")
    uploaded_json = st.file_uploader("选择 policy_data.json", type="json")
    json_data = None
    if uploaded_json:
        try:
            json_data = json.load(uploaded_json)
            data_errors = validate_json_data(json_data)
            if data_errors:
                st.warning("JSON 数据检查发现问题：")
                for err in data_errors:
                    st.write(f"- {err}")
            else:
                st.success("✅ 数据加载成功，结构检查通过")
        except Exception:
            st.error("❌ JSON 格式错误")

with col2:
    st.subheader("📄 2. 上传模板 (Word)")
    uploaded_template = st.file_uploader("选择 模板.docx", type="docx")
    if uploaded_template:
        st.success("✅ 模板加载成功")

st.markdown("---")

can_generate = uploaded_json and uploaded_template and json_data is not None
if st.button("🚀 开始生成报告", disabled=not can_generate):
    start = time.time()
    status = "success"
    try:
        preview_doc = Document(uploaded_template)
        template_check = validate_template(preview_doc, section_config, expected_placeholder_count=31)

        if template_check["missing_titles"]:
            st.warning("模板缺少以下章节标题：")
            for title in template_check["missing_titles"]:
                st.write(f"- {title}")

        if not template_check["chart_anchor_ok"]:
            st.warning(f"未找到图表锚点：{section_config.get('chart_anchor')}")

        if not template_check["placeholder_ok"]:
            st.warning(
                f"占位符数量不足：模板 {template_check['placeholder_count']} 个，"
                f"预期至少 {template_check['expected_placeholder_count']} 个"
            )

        uploaded_template.seek(0)
        with st.spinner("正在处理：模板预检 -> 替换数据 -> 绘图 -> RAG检索 -> AI解析 -> 排版 ..."):
            result_doc = process_document(
                template_file=uploaded_template,
                json_data=json_data,
                api_key=api_key,
                section_config=section_config,
                rag_chunks=rag_chunks,
                rag_top_k=rag_top_k,
                include_rag_snippets=include_rag_snippets,
            )

        st.balloons()
        st.success("🎉 报告生成成功！")
        st.download_button(
            label="📥 下载 Word 文档",
            data=result_doc,
            file_name="政策资讯专刊_自动生成.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as err:
        status = "failed"
        LOGGER.exception("报告生成失败")
        st.error(f"生成失败: {str(err)}")
    finally:
        duration = round(time.time() - start, 2)
        history_record = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": status,
            "duration_s": duration,
            "rag_enabled": bool(enable_rag),
            "rag_chunks": len(rag_chunks),
            "template": getattr(uploaded_template, "name", "unknown"),
        }
        append_history(history_record)
        LOGGER.info("generation=%s", json.dumps(history_record, ensure_ascii=False))
