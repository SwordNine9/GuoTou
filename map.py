import os
import json
import requests
import matplotlib
matplotlib.use("Agg")  # 无图形界面环境也能画图
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from openai import OpenAI

# ========== 配置部分 ==========
# DeepSeek API配置（需要替换为实际API URL和密钥）
DEEPSEEK_API_URL = "https://api.deepseek.com"
DEEPSEEK_API_KEY = "YOUR_DEEPSEEK_API_KEY"  # TODO: 在这里填你的 Key，或用环境变量

# 样例模板和输出文件路径
template_path = "国投电力业务发展部政策资讯专刊（模板）.docx"
output_docx = "政策资讯专刊_自动生成.docx"
output_pdf = "政策资讯专刊_自动生成.pdf"

# ========== 从JSON文件读取数据 ==========
def load_json_data(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)

# 读取JSON文件
data = load_json_data("policy_data.json")
if data:
    print("Data loaded successfully:", data)
else:
    print("Failed to load data.")

# 重要数据
important_data_values = data.get("important_data_values", {})
report_date = important_data_values.get("report_date", "")

# 政府文件政策数据
gov_policies = data.get("gov_policies", [])

# 能源新政数据
energy_new_policies = data.get("energy_new_policies", {})

# 参考资料
references = data.get("references", [])

# ========== 文档生成部分 ==========
try:
    # 尝试加载模板
    doc = Document(template_path)
    print("Template loaded successfully.")
except Exception as e:
    print(f"Error loading template: {e}")
    doc = None

# 工具函数：替换段落中下一个“**”占位符为指定值（保留样式）
def replace_next_placeholder(paragraph, value):
    """
    在paragraph中查找下一个“**”并替换为value，替换一次后返回True；如果未找到则返回False。
    """
    for run in paragraph.runs:
        if "**" in run.text:
            # 仅替换当前 run 中的第一个占位符
            run.text = run.text.replace("**", str(value), 1)
            return True
    return False

# ========== 图表生成相关函数（饼状图） ==========

# 生成水电、火电、核电、太阳能、风电的装机容量饼状图
def generate_capacity_pie_chart(important_values, chart_path):
    import matplotlib.pyplot as plt

    # 防止负号显示成方块（习惯性加上）
    plt.rcParams['axes.unicode_minus'] = False

    # 图上的标签改成英文
    categories = ["Hydro", "Thermal", "Nuclear", "Solar", "Wind"]
    keys = [
        "hydro_capacity",
        "thermal_capacity",
        "nuclear_capacity",
        "solar_capacity",
        "wind_capacity",
    ]

    values = []
    for k in keys:
        v = important_values.get(k, 0) or 0
        try:
            values.append(float(v))
        except ValueError:
            values.append(0.0)

    # 防止全为 0 导致饼图报错
    if all(v == 0 for v in values):
        print("容量全为 0，无法生成饼状图。")
        return

    plt.figure(figsize=(6, 6))
    wedges, texts, autotexts = plt.pie(
        values,
        labels=categories,        # 扇区外的英文标签
        autopct="%.1f%%",         # 扇区内的百分比
        startangle=90,
        labeldistance=1.1,
        pctdistance=0.75
    )

    # 图标题也改成英文，避免中文字体问题
    plt.title("Power Installed Capacity Structure (by Energy Type)")

    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()
    print(f"装机容量饼状图已生成：{chart_path}")



# 将图表插入到文档（当前逻辑：插入到文档末尾）
def insert_chart_into_doc(doc, chart_path):
    if not doc:
        return
    doc.add_paragraph("图 电力装机情况（自动生成饼状图）")
    doc.add_picture(chart_path, width=Inches(5.0))
    print("图表已插入到文档中。")

# ========== 填充模版中的数值 ==========
if doc:
    try:
        imp_values_iter = iter([
            report_date,
            important_data_values.get("total_capacity", ""),
            important_data_values.get("total_capacity_yoy", ""),
            important_data_values.get("hydro_capacity", ""),
            important_data_values.get("hydro_yoy", ""),
            important_data_values.get("thermal_capacity", ""),
            important_data_values.get("thermal_yoy", ""),
            important_data_values.get("nuclear_capacity", ""),
            important_data_values.get("nuclear_yoy", ""),
            important_data_values.get("solar_capacity", ""),
            important_data_values.get("solar_yoy", ""),
            important_data_values.get("wind_capacity", ""),
            important_data_values.get("wind_yoy", ""),
            report_date,
            important_data_values.get("market_trade", ""),
            important_data_values.get("market_trade_yoy", ""),
            important_data_values.get("provincial_trade", ""),
            important_data_values.get("provincial_trade_yoy", ""),
            important_data_values.get("cross_region_trade", ""),
            important_data_values.get("cross_region_trade_yoy", ""),
            important_data_values.get("green_trade", ""),
            important_data_values.get("green_trade_yoy", ""),
            report_date,
            important_data_values.get("society_usage", ""),
            important_data_values.get("society_usage_yoy", ""),
            important_data_values.get("primary_usage", ""),
            important_data_values.get("primary_usage_yoy", ""),
            important_data_values.get("secondary_usage", ""),
            important_data_values.get("secondary_usage_yoy", ""),
            important_data_values.get("tertiary_usage", ""),
            important_data_values.get("tertiary_usage_yoy", "")
        ])

        # 遍历段落并替换占位符
        for paragraph in doc.paragraphs:
            while "**" in paragraph.text:
                try:
                    value = next(imp_values_iter)
                except StopIteration:
                    print("All placeholders replaced successfully.")
                    break  # 完成替换后退出循环
                replace_next_placeholder(paragraph, value)

        print("Template filled successfully.")

    except Exception as e:
        print(f"Error during document generation: {e}")

# ========== 政府文件部分 ==========
def insert_government_policy(gov_policies):
    if not doc:
        return
    for policy in gov_policies:
        title = policy['title']
        agency = policy['agency']
        date = policy['date']
        summary = policy['summary']

        doc.add_paragraph(f"政策名称：{title}")
        doc.add_paragraph(f"发布机构：{agency}")
        doc.add_paragraph(f"发布时间：{date}")
        doc.add_paragraph(f"政策摘要：{summary}")

        analysis_text = generate_policy_analysis(summary)

        doc.add_paragraph(f"政策解析：{analysis_text}")
        doc.add_paragraph("")

# ========== 能源新政部分 ==========
def insert_energy_new_policies(energy_new_policies):
    if not doc:
        return
    for topic, policies in energy_new_policies.items():
        doc.add_paragraph(f"{topic}")
        for policy in policies:
            doc.add_paragraph(f"区域：{policy['region']}")
            doc.add_paragraph(f"政策标题：{policy['title']}")
            doc.add_paragraph(f"发布日期：{policy['date']}")
            doc.add_paragraph(f"政策摘要：{policy['summary']}")
            analysis_text = generate_policy_analysis(policy['summary'])

            # 插入政策分析
            doc.add_paragraph(f"政策解析：{analysis_text}")
            doc.add_paragraph("")

# ========== DeepSeek接口：生成政策解析内容 ==========
def generate_policy_analysis(policy_summary):
    try:
        # 使用 OpenAI SDK 格式发送 DeepSeek API 请求
        client = OpenAI(
            api_key=DEEPSEEK_API_KEY,  # 建议从环境变量获取API密钥
            base_url=DEEPSEEK_API_URL
        )
        response = client.chat.completions.create(
            model="deepseek-chat",  # 模型名称，根据 DeepSeek 文档设置
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是一位电力政策分析专家，根据我为你提供的政策摘要生成一份政策分析，"
                        "请基于以下政策摘要生成一个简洁的政策分析。分析内容应包括："
                        "1. 主要政策要点 2. 可能的影响和挑战 3. 相关建议和应对策略，"
                        "所有的内容都在一个段落里"
                    )
                },
                {"role": "user", "content": policy_summary}  # 传入政策摘要
            ],
            stream=False  # 禁用流式响应
        )

        # 打印 API 返回的数据
        print(f"DeepSeek API response data: {response}")
        
        # 获取并返回分析内容
        return response.choices[0].message.content  # 返回分析内容
    except Exception as e:
        print(f"Error during API request: {e}")
        return "API调用失败，无法生成分析"
    
# ========== 组装文档 ==========
# 插入政策数据
insert_government_policy(gov_policies)

# 插入能源新政部分
insert_energy_new_policies(energy_new_policies)

# 生成并插入饼状图
chart_path = "装机容量结构饼图.png"
generate_capacity_pie_chart(important_data_values, chart_path)
insert_chart_into_doc(doc, chart_path)

# 保存生成的Word文档
if doc:
    try:
        doc.save(output_docx)
        print("Document saved successfully.")
    except Exception as e:
        print(f"Error saving document: {e}")

# 可选：将生成的Word转换为PDF
# from docx2pdf import convert
# try:
#     convert(output_docx, output_pdf)
# except Exception as e:
#     print("PDF转换失败：", e)
